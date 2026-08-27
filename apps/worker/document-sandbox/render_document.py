import argparse
import base64
import datetime
import io
import json
import re
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt
from PIL import Image
from pypdf import PdfReader
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas


def read_source(path: Path) -> tuple[list[dict[str, object]], dict[str, bytes]]:
    value = json.loads(path.read_text(encoding="utf-8"))
    pages = value.get("pages") if isinstance(value, dict) else None
    if not isinstance(pages, list) or not 1 <= len(pages) <= 20:
        raise ValueError("source must contain between 1 and 20 pages")
    for page in pages:
        if not isinstance(page, dict):
            raise ValueError("each page must be an object")
        title = page.get("title")
        lines = page.get("lines")
        if not isinstance(title, str) or not 1 <= len(title) <= 80:
            raise ValueError("each page title must contain 1 to 80 characters")
        if not isinstance(lines, list) or len(lines) > 30:
            raise ValueError("each page must contain at most 30 lines")
        if any(not isinstance(line, str) or len(line) > 80 for line in lines):
            raise ValueError("each line must contain at most 80 characters")
        visual_content_id = page.get("visualContentId")
        if visual_content_id is not None and (
            not isinstance(visual_content_id, str) or len(lines) > 16
        ):
            raise ValueError("a visual page must reference one identity and at most 16 lines")
    encoded_visuals = value.get("supportingVisuals", [])
    if not isinstance(encoded_visuals, list) or len(encoded_visuals) > 20:
        raise ValueError("supporting visuals must be a bounded list")
    visuals: dict[str, bytes] = {}
    total_visual_bytes = 0
    for encoded_visual in encoded_visuals:
        if not isinstance(encoded_visual, dict):
            raise ValueError("supporting visual must be an object")
        content_id = encoded_visual.get("contentId")
        body = encoded_visual.get("base64")
        if not isinstance(content_id, str) or not isinstance(body, str):
            raise ValueError("supporting visual identity and body are required")
        if content_id in visuals:
            raise ValueError("supporting visual identities must be unique")
        decoded = base64.b64decode(body, validate=True)
        total_visual_bytes += len(decoded)
        if total_visual_bytes > 25_000_000:
            raise ValueError("supporting visuals exceed the immutable input limit")
        with Image.open(io.BytesIO(decoded)) as image:
            image.verify()
        visuals[content_id] = decoded
    referenced = {
        visual_content_id
        for page in pages
        if isinstance((visual_content_id := page.get("visualContentId")), str)
    }
    if referenced != set(visuals):
        raise ValueError("supporting visuals must exactly match page references")
    return pages, visuals


def render_pdf(pages: list[dict[str, object]], visuals: dict[str, bytes], output: Path) -> int:
    document = canvas.Canvas(str(output), pagesize=LETTER, invariant=1, pageCompression=1)
    for page in pages:
        document.setFont("Helvetica-Bold", 16)
        document.drawString(54, 738, page["title"])
        document.setFont("Helvetica", 11)
        y = 708
        for line in page["lines"]:
            document.drawString(54, y, line)
            y -= 20
        visual_content_id = page.get("visualContentId")
        if isinstance(visual_content_id, str):
            document.drawImage(
                ImageReader(io.BytesIO(visuals[visual_content_id])),
                324,
                72,
                width=234,
                height=234,
                preserveAspectRatio=True,
                anchor="c",
            )
        document.showPage()
    document.save()
    return len(PdfReader(output).pages)


def render_docx(pages: list[dict[str, object]], visuals: dict[str, bytes], output: Path) -> int:
    document = Document()
    for section in document.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    normal = document.styles["Normal"]
    normal.font.name = "Liberation Mono"
    normal.font.size = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    heading = document.styles["Heading 1"]
    heading.font.name = "Liberation Mono"
    heading.font.size = Pt(12)
    heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading.paragraph_format.line_spacing = Pt(14)
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(4)
    fixed_time = datetime.datetime(2000, 1, 1, tzinfo=datetime.timezone.utc)
    document.core_properties.created = fixed_time
    document.core_properties.modified = fixed_time
    for page_index, page in enumerate(pages):
        document.add_heading(page["title"], level=1)
        for line in page["lines"]:
            document.add_paragraph(line)
        visual_content_id = page.get("visualContentId")
        if isinstance(visual_content_id, str):
            visual = visuals[visual_content_id]
            with Image.open(io.BytesIO(visual)) as image:
                scale = min(3 / image.width, 3 / image.height)
                width = Inches(image.width * scale)
                height = Inches(image.height * scale)
            document.add_picture(io.BytesIO(visual), width=width, height=height)
        if page_index + 1 < len(pages):
            document.add_page_break()
    temporary = output.with_suffix(".raw.docx")
    document.save(temporary)
    with zipfile.ZipFile(temporary, "r") as source, zipfile.ZipFile(
        output, "w", zipfile.ZIP_DEFLATED
    ) as target:
        for item in sorted(source.infolist(), key=lambda candidate: candidate.filename):
            content = source.read(item.filename)
            if item.filename == "docProps/app.xml":
                text = content.decode("utf-8")
                text = re.sub(r"<Pages>.*?</Pages>", f"<Pages>{len(pages)}</Pages>", text)
                content = text.encode("utf-8")
            deterministic = zipfile.ZipInfo(item.filename, (2000, 1, 1, 0, 0, 0))
            deterministic.compress_type = zipfile.ZIP_DEFLATED
            deterministic.external_attr = item.external_attr
            target.writestr(deterministic, content)
    temporary.unlink()
    with tempfile.TemporaryDirectory() as rendered_directory:
        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                rendered_directory,
                str(output),
            ],
            check=True,
            capture_output=True,
            timeout=30,
        )
        rendered_pages = len(PdfReader(Path(rendered_directory) / f"{output.stem}.pdf").pages)
    if rendered_pages != len(pages):
        raise ValueError("DOCX layout does not preserve one rendered page per source page")
    return rendered_pages


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--format", choices=("pdf", "docx"), required=True)
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    arguments = parser.parse_args()
    pages, visuals = read_source(arguments.input)
    if arguments.format == "pdf":
        rendered_pages = render_pdf(pages, visuals, arguments.output)
    else:
        rendered_pages = render_docx(pages, visuals, arguments.output)
    print(json.dumps({"renderedPageCount": rendered_pages}))


if __name__ == "__main__":
    main()
