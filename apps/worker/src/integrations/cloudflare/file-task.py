import csv
import io
import json
import sys
import zipfile

from PIL import Image
from docx import Document
from pypdf import PdfReader


class ContentLimit(Exception):
    pass


class MaliciousContent(Exception):
    pass


def bounded_text(text, limit):
    encoded = text.encode("utf-8")
    if len(encoded) > limit:
        raise ContentLimit("normalized text exceeds its byte limit")
    return text


def normalize(config):
    media_type = config["mediaType"]
    limits = config["limits"]
    source_path = "/workspace/source.bin"

    if media_type in ("text/plain", "text/csv"):
        with open(source_path, "r", encoding="utf-8", errors="strict", newline="") as source:
            text = source.read()
        if media_type == "text/csv":
            rows = sum(1 for _ in csv.reader(io.StringIO(text)))
            if rows > limits["maximumCsvRows"]:
                raise ContentLimit("CSV row count exceeds its limit")
        parser = "python-stdlib-v1"
    elif media_type == "application/pdf":
        reader = PdfReader(source_path, strict=True)
        if reader.is_encrypted:
            raise MaliciousContent("encrypted PDFs are not accepted")
        if len(reader.pages) > limits["maximumPdfPages"]:
            raise ContentLimit("PDF page count exceeds its limit")
        text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        parser = "pypdf-5.9.0"
    elif media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        with zipfile.ZipFile(source_path) as archive:
            entries = archive.infolist()
            if len(entries) > limits["maximumOfficeEntries"]:
                raise ContentLimit("DOCX entry count exceeds its limit")
            compressed = sum(entry.compress_size for entry in entries)
            expanded = sum(entry.file_size for entry in entries)
            if expanded > limits["maximumNormalizedTextBytes"] * 20:
                raise ContentLimit("DOCX expanded size exceeds its limit")
            if compressed > 0 and expanded > compressed * 100:
                raise MaliciousContent("DOCX expansion ratio is unsafe")
        document = Document(source_path)
        blocks = [paragraph.text for paragraph in document.paragraphs]
        blocks.extend("\t".join(cell.text for cell in row.cells) for table in document.tables for row in table.rows)
        text = "\n".join(blocks)
        parser = "python-docx-1.2.0"
    else:
        Image.MAX_IMAGE_PIXELS = limits["maximumImagePixels"]
        with Image.open(source_path) as image:
            image.verify()
        with Image.open(source_path) as image:
            if image.width * image.height > limits["maximumImagePixels"]:
                raise ContentLimit("image pixel count exceeds its limit")
            text = f"Image: format={image.format}; width={image.width}; height={image.height}; mode={image.mode}"
        parser = "pillow-12.3.0"

    return {
        "normalizedText": bounded_text(text, limits["maximumNormalizedTextBytes"]),
        "parser": parser,
    }


def analyze(config):
    text = config["normalizedText"]
    prompt = config["prompt"].strip()
    words = text.split()
    preview = text[:2000]
    result = f"Request: {prompt}\nCharacters: {len(text)}\nWords: {len(words)}\nContent preview:\n{preview}"
    return {"resultText": bounded_text(result, 2000000)}


def main():
    with open("/workspace/input.json", "r", encoding="utf-8") as source:
        config = json.load(source)
    operation = config["operation"]
    result = normalize(config) if operation == "normalize" else analyze(config)
    output = {"ok": True, **result}
    with open("/workspace/result.json", "w", encoding="utf-8") as target:
        json.dump(output, target, ensure_ascii=False)


try:
    main()
except ContentLimit as error:
    with open("/workspace/result.json", "w", encoding="utf-8") as target:
        json.dump({"ok": False, "reason": "content_limit", "message": str(error)}, target)
except (Image.DecompressionBombError, MaliciousContent) as error:
    with open("/workspace/result.json", "w", encoding="utf-8") as target:
        json.dump({"ok": False, "reason": "malicious", "message": str(error)}, target)
except Exception as error:
    with open("/workspace/result.json", "w", encoding="utf-8") as target:
        json.dump({"ok": False, "reason": "parser_failure", "message": str(error)}, target)
    sys.exit(1)
